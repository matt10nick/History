from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import math 
import imageeditor

def generate_document(list_of_files, image_directory, output_path):
    # Create a new Document
    doc = Document()

    pages = math.ceil(len(list_of_files) / 4)
    file_index = 0

    for page in range(pages):

        images_left = len(list_of_files) - file_index
        images_on_page = min(images_left, 4)

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        for index in range(images_on_page):

            file = list_of_files[page*4 + index]
            path = image_directory + "/" + file

            tmp_file_name = "./tmp/rotated_image" + str(index) + ".jpg"        
            imageeditor.rotate_image(path, tmp_file_name, 90)

            run.add_picture(tmp_file_name, width=Inches(4))

        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        doc.add_page_break()

        paragraph = doc.add_paragraph()
        run = paragraph.add_run()

        for index in range(images_on_page):

            tmp_file_name = "./tmp/rotated_image" + str(index) + ".jpg"        
            imageeditor.rotate_image(tmp_file_name, tmp_file_name, 180)

            run.add_picture(tmp_file_name, width=Inches(4))
        
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


        file_index += images_on_page
    
    # Save the document
    doc.save(output_path)